import PyPDF2
from docx import *
from django.shortcuts import render
from core.algorithm import main, fileSimilarity

# Custom error pages


def error_404(request, exception):
    return render(request, '404.html')


def error_500(request):
    return render(request, '500.html')


def landingPage(request):
    return render(request, 'landingPage.html')


def file(request):
    return render(request, 'file.html')


def text(request):
    return render(request, 'text.html')


def compare(request):
    return render(request, 'compare.html')

def docs(request):
    return render(request, 'docs.html')


# this function is used to find palagrism in text and return the result to the user in the form of percentage and link
def result(request):
    print(request.POST['q'])

    if request.POST['q']:
        percent, link = main.findSimilarity(request.POST['q'])
        percent = round(percent, 2)
    print("Output: ", percent, link)
    percent = round(percent, 2)
    return render(request, 'resultPage.html', {'link': link, 'percent': percent})


# this function is used to find palagrism in file and return the result to the user in the form of percentage and link
def fileTestResult(request):
    value = ''
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object
        pdfFileObj = open(request.FILES['docfile'], 'rb')

        # creating a pdf reader object
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        # printing number of pages in pdf file
        print(pdfReader.numPages)

        # creating a page object
        pageObj = pdfReader.getPage(0)

        # extracting text from page
        print(pageObj.extractText())

        # closing the pdf file object
        pdfFileObj.close()

    percent, link = main.findSimilarity(value)
    percent = round(percent, 2)
    print("Output: ", percent, link)
    return render(request, 'resultPage.html', {'link': link, 'percent': percent})


# this function is used to find the similarity between two texts
def textSimilarityAnalysis(request):
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '':
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(
            request.POST['q1'], request.POST['q2'])
    result = round(result, 2)
    print("Output: ", result)
    return render(request, 'compare.html', {'result': result})


# this function is used to find the similarity between two files
def fileSimilarityAnalysis(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1, value2)
    result = round(result, 2)
    print("Output: ", result)
    return render(request, 'compare.html', {'result': result})
